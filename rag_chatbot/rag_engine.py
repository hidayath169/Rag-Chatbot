"""
RAG Engine — Document processing, embedding, and retrieval pipeline.
Uses LangChain + FAISS + Google Gemini for resume Q&A.
"""

import os
import uuid
import json
import shutil
from datetime import datetime
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.messages import HumanMessage, AIMessage
from pypdf import PdfReader
from docx import Document as DocxDocument


# ── Paths ──────────────────────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent
FAISS_INDEX_DIR = BASE_DIR / "faiss_index"
UPLOADS_DIR = BASE_DIR / "uploads"
METADATA_FILE = BASE_DIR / "document_metadata.json"

for d in [FAISS_INDEX_DIR, UPLOADS_DIR]:
    d.mkdir(exist_ok=True)


# ── Document Processor ────────────────────────────────────────────────────────

class DocumentProcessor:
    """Handles PDF/DOCX parsing and text chunking."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def load_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        reader = PdfReader(file_path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    def load_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def load_file(self, file_path: str) -> str:
        """Load text from a supported file format."""
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self.load_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self.load_docx(file_path)
        elif ext == ".txt":
            return Path(file_path).read_text(encoding="utf-8", errors="ignore")
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def process_file(self, file_path: str, filename: str) -> list[Document]:
        """Parse a file and split into LangChain Documents with metadata."""
        text = self.load_file(file_path)
        if not text.strip():
            return []

        doc_id = str(uuid.uuid4())[:8]
        metadata = {
            "source": filename,
            "doc_id": doc_id,
            "uploaded_at": datetime.now().isoformat(),
        }

        # Create a single Document then split
        full_doc = Document(page_content=text, metadata=metadata)
        chunks = self.splitter.split_documents([full_doc])

        # Add chunk index to each
        for i, chunk in enumerate(chunks):
            chunk.metadata["chunk_index"] = i
            chunk.metadata["total_chunks"] = len(chunks)

        return chunks


# ── RAG Engine ─────────────────────────────────────────────────────────────────

class RAGEngine:
    """Manages the FAISS vector store and LLM-powered Q&A."""

    def __init__(self, api_key: str | None = None):
        self.api_key = api_key or os.getenv("GOOGLE_API_KEY", "")
        self.processor = DocumentProcessor()
        self.vectorstore: FAISS | None = None
        self.doc_metadata: dict = {}  # filename -> {doc_id, uploaded_at, chunks}
        self._load_metadata()
        self._load_index()

    # ── Private helpers ────────────────────────────────────────────────────

    def _get_embeddings(self):
        return GoogleGenerativeAIEmbeddings(
            model="models/gemini-embedding-2",
            google_api_key=self.api_key,
        )

    def _get_llm(self):
        return ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=self.api_key,
            temperature=0.3,
        )

    def _save_metadata(self):
        with open(METADATA_FILE, "w") as f:
            json.dump(self.doc_metadata, f, indent=2)

    def _load_metadata(self):
        if METADATA_FILE.exists():
            with open(METADATA_FILE, "r") as f:
                self.doc_metadata = json.load(f)
        else:
            self.doc_metadata = {}

    def _save_index(self):
        if self.vectorstore:
            self.vectorstore.save_local(str(FAISS_INDEX_DIR))

    def _load_index(self):
        index_file = FAISS_INDEX_DIR / "index.faiss"
        if index_file.exists() and self.api_key:
            try:
                self.vectorstore = FAISS.load_local(
                    str(FAISS_INDEX_DIR),
                    self._get_embeddings(),
                    allow_dangerous_deserialization=True,
                )
            except Exception as e:
                print(f"Warning: Could not load FAISS index: {e}")
                self.vectorstore = None

    # ── Public API ─────────────────────────────────────────────────────────

    def set_api_key(self, key: str):
        """Update the API key and reload the index."""
        self.api_key = key
        os.environ["GOOGLE_API_KEY"] = key
        self._load_index()

    def add_documents(self, file_paths: list[tuple[str, str]]) -> dict:
        """
        Process and index multiple files.
        file_paths: list of (temp_file_path, original_filename)
        Returns: {added: int, errors: list}
        """
        all_chunks = []
        errors = []

        for file_path, filename in file_paths:
            try:
                chunks = self.processor.process_file(file_path, filename)
                if not chunks:
                    errors.append(f"{filename}: No text extracted")
                    continue

                all_chunks.extend(chunks)
                self.doc_metadata[filename] = {
                    "doc_id": chunks[0].metadata["doc_id"],
                    "uploaded_at": chunks[0].metadata["uploaded_at"],
                    "chunks": len(chunks),
                }
            except Exception as e:
                errors.append(f"{filename}: {str(e)}")

        if all_chunks:
            embeddings = self._get_embeddings()
            if self.vectorstore is None:
                self.vectorstore = FAISS.from_documents(all_chunks, embeddings)
            else:
                new_store = FAISS.from_documents(all_chunks, embeddings)
                self.vectorstore.merge_from(new_store)

            self._save_index()
            self._save_metadata()

        return {"added": len(file_paths) - len(errors), "errors": errors}

    def query(self, question: str, chat_history: list[dict] | None = None) -> str:
        """
        Answer a question using the RAG pipeline.
        chat_history: list of {role: 'user'|'assistant', content: str}
        """
        if not self.api_key:
            return "⚠️ Please set your Google API key in the Settings panel before asking questions."

        if not self.vectorstore:
            return "📂 No resumes uploaded yet. Upload some resumes first, then ask me questions about them!"

        # Retrieve relevant chunks
        retriever = self.vectorstore.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 6},
        )
        relevant_docs = retriever.invoke(question)

        if not relevant_docs:
            return "I couldn't find any relevant information in the uploaded resumes for your question."

        # Build context from retrieved chunks
        context_parts = []
        for doc in relevant_docs:
            source = doc.metadata.get("source", "Unknown")
            context_parts.append(f"[Resume: {source}]\n{doc.page_content}")
        context = "\n\n---\n\n".join(context_parts)

        # Build prompt
        prompt = ChatPromptTemplate.from_messages([
            (
                "system",
                """You are an intelligent Resume Analysis Assistant. You help users find information from uploaded resumes.

INSTRUCTIONS:
- Answer questions based ONLY on the resume data provided in the context below.
- Always mention which resume/candidate the information comes from.
- If the information is not found in any resume, say so clearly — do NOT make up information.
- Format your answers clearly with bullet points or tables when appropriate.
- Use markdown formatting for better readability.
- When comparing candidates, create a structured comparison.

RESUME CONTEXT:
{context}"""
            ),
            MessagesPlaceholder(variable_name="history"),
            ("human", "{question}"),
        ])

        # Convert chat history
        history_messages = []
        if chat_history:
            for msg in chat_history[-10:]:  # Keep last 10 messages for context
                if msg["role"] == "user":
                    history_messages.append(HumanMessage(content=msg["content"]))
                elif msg["role"] == "assistant":
                    history_messages.append(AIMessage(content=msg["content"]))

        # Generate response
        llm = self._get_llm()
        chain = prompt | llm
        response = chain.invoke({
            "context": context,
            "history": history_messages,
            "question": question,
        })

        return response.content

    def list_documents(self) -> list[dict]:
        """Return metadata for all indexed documents."""
        docs = []
        for filename, meta in self.doc_metadata.items():
            docs.append({
                "filename": filename,
                "doc_id": meta["doc_id"],
                "uploaded_at": meta["uploaded_at"],
                "chunks": meta["chunks"],
            })
        return sorted(docs, key=lambda x: x["uploaded_at"], reverse=True)

    def delete_document(self, filename: str) -> bool:
        """
        Remove a document from the metadata and rebuild the index.
        (FAISS doesn't support deletion, so we rebuild without the deleted doc's chunks.)
        """
        if filename not in self.doc_metadata:
            return False

        target_doc_id = self.doc_metadata[filename]["doc_id"]
        del self.doc_metadata[filename]
        self._save_metadata()

        # Rebuild index without the deleted document
        if self.vectorstore and self.doc_metadata:
            # Get all documents from the store
            all_docs = []
            docstore = self.vectorstore.docstore
            index_to_id = self.vectorstore.index_to_docstore_id

            for idx, doc_id in index_to_id.items():
                doc = docstore.search(doc_id)
                if hasattr(doc, 'metadata') and doc.metadata.get("doc_id") != target_doc_id:
                    all_docs.append(doc)

            if all_docs:
                embeddings = self._get_embeddings()
                self.vectorstore = FAISS.from_documents(all_docs, embeddings)
                self._save_index()
            else:
                self.vectorstore = None
                # Clean up index files
                if FAISS_INDEX_DIR.exists():
                    shutil.rmtree(FAISS_INDEX_DIR)
                    FAISS_INDEX_DIR.mkdir(exist_ok=True)
        else:
            self.vectorstore = None
            if FAISS_INDEX_DIR.exists():
                shutil.rmtree(FAISS_INDEX_DIR)
                FAISS_INDEX_DIR.mkdir(exist_ok=True)

        return True

    def get_stats(self) -> dict:
        """Return statistics about the vector store."""
        total_chunks = sum(m["chunks"] for m in self.doc_metadata.values())
        return {
            "total_documents": len(self.doc_metadata),
            "total_chunks": total_chunks,
            "index_loaded": self.vectorstore is not None,
        }
